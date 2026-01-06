"""
Script tạo báo cáo kiểm tra dự án Flutter
Chạy: python create_report.py
Yêu cầu: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Đặt border cho ô trong bảng"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('left', 'top', 'right', 'bottom'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:color'), '000000')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_report():
    doc = Document()
    
    # Tiêu đề chính
    title = doc.add_heading('BÁO CÁO KIỂM TRA DỰ ÁN FLUTTER', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin môn học
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Môn học: ').bold = True
    info.add_run('Lập trình di động')
    info.add_run('\nNgười kiểm tra: ').bold = True
    info.add_run('[Họ và tên của bạn]')
    info.add_run('\nNgày kiểm tra: ').bold = True
    info.add_run('[Ngày/Tháng/Năm]')
    
    doc.add_paragraph()
    
    # ========== NHÓM 1 ==========
    groups_data = [
        {
            "group_name": "Nhóm 1",
            "members": ["Nguyễn Văn A", "Trần Thị B", "Lê Văn C"],
            "errors": [
                {
                    "name": "Lỗi 1: [Tên lỗi]",
                    "description": "[Mô tả chi tiết lỗi]",
                    "image_placeholder": "[Chèn ảnh minh chứng tại đây]"
                },
                {
                    "name": "Lỗi 2: [Tên lỗi]",
                    "description": "[Mô tả chi tiết lỗi]",
                    "image_placeholder": "[Chèn ảnh minh chứng tại đây]"
                }
            ]
        },
        {
            "group_name": "Nhóm 2",
            "members": ["[Thành viên 1]", "[Thành viên 2]", "[Thành viên 3]"],
            "errors": [
                {
                    "name": "Lỗi 1: [Tên lỗi]",
                    "description": "[Mô tả chi tiết lỗi]",
                    "image_placeholder": "[Chèn ảnh minh chứng tại đây]"
                }
            ]
        },
        {
            "group_name": "Nhóm 3",
            "members": ["[Thành viên 1]", "[Thành viên 2]"],
            "errors": [
                {
                    "name": "Lỗi 1: [Tên lỗi]",
                    "description": "[Mô tả chi tiết lỗi]",
                    "image_placeholder": "[Chèn ảnh minh chứng tại đây]"
                }
            ]
        }
    ]
    
    for group in groups_data:
        # Tiêu đề nhóm
        group_heading = doc.add_heading(group["group_name"], level=1)
        group_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Bảng thông tin nhóm
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Row 1: Tên nhóm
        table.cell(0, 0).text = "Tên nhóm"
        table.cell(0, 1).text = group["group_name"]
        
        # Row 2: Thành viên
        table.cell(1, 0).text = "Thành viên"
        members_text = "\n".join([f"• {member}" for member in group["members"]])
        table.cell(1, 1).text = members_text
        
        # Bold headers
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Danh sách lỗi
        errors_heading = doc.add_heading("Danh sách lỗi phát hiện:", level=2)
        
        for i, error in enumerate(group["errors"], 1):
            # Tên lỗi
            error_para = doc.add_paragraph()
            error_para.add_run(f"{i}. {error['name']}").bold = True
            
            # Mô tả lỗi
            desc_para = doc.add_paragraph()
            desc_para.add_run("Mô tả: ").bold = True
            desc_para.add_run(error['description'])
            
            # Placeholder cho ảnh
            img_para = doc.add_paragraph()
            img_para.add_run("Ảnh minh chứng:").bold = True
            
            # Tạo khung cho ảnh
            img_box = doc.add_paragraph()
            img_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_box.add_run()
            run.add_text("┌─────────────────────────────────────────┐")
            img_box.add_run("\n│                                         │")
            img_box.add_run("\n│     " + error['image_placeholder'] + "      │")
            img_box.add_run("\n│                                         │")
            img_box.add_run("\n└─────────────────────────────────────────┘")
            
            doc.add_paragraph()
        
        # Đường kẻ phân cách giữa các nhóm
        doc.add_paragraph("─" * 60)
        doc.add_paragraph()
    
    # Phần tổng kết
    summary = doc.add_heading("TỔNG KẾT", level=1)
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    summary_table = doc.add_table(rows=4, cols=3)
    summary_table.style = 'Table Grid'
    
    # Header
    summary_table.cell(0, 0).text = "Nhóm"
    summary_table.cell(0, 1).text = "Số lỗi"
    summary_table.cell(0, 2).text = "Đánh giá"
    
    # Data rows
    summary_table.cell(1, 0).text = "Nhóm 1"
    summary_table.cell(1, 1).text = "[Số lỗi]"
    summary_table.cell(1, 2).text = "[Tốt/Khá/Trung bình/Yếu]"
    
    summary_table.cell(2, 0).text = "Nhóm 2"
    summary_table.cell(2, 1).text = "[Số lỗi]"
    summary_table.cell(2, 2).text = "[Tốt/Khá/Trung bình/Yếu]"
    
    summary_table.cell(3, 0).text = "Nhóm 3"
    summary_table.cell(3, 1).text = "[Số lỗi]"
    summary_table.cell(3, 2).text = "[Tốt/Khá/Trung bình/Yếu]"
    
    # Bold header row
    for cell in summary_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Chữ ký
    doc.add_paragraph()
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.add_run("Người kiểm tra\n\n\n\n")
    signature.add_run("[Ký tên]")
    
    # Lưu file
    filename = "Bao_Cao_Kiem_Tra_Du_An_Flutter.docx"
    doc.save(filename)
    print(f"✅ Đã tạo file báo cáo: {filename}")
    print(f"📁 Đường dẫn: {filename}")
    print("\n📝 Hướng dẫn:")
    print("1. Mở file Word và chỉnh sửa thông tin các nhóm")
    print("2. Thay thế [Chèn ảnh minh chứng tại đây] bằng ảnh thực tế")
    print("3. Để chèn ảnh: Insert -> Pictures -> Chọn ảnh")

if __name__ == "__main__":
    create_report()
