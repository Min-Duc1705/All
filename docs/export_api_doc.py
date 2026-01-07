"""
Script để xuất API Documentation từ Markdown sang Word (.docx)
Sử dụng thư viện python-docx và markdown

Chạy: python export_api_doc.py
Kết quả: API_Documentation.docx
"""

import subprocess
import sys

# Cài đặt các thư viện cần thiết
def install_packages():
    packages = ['python-docx', 'markdown', 'beautifulsoup4', 'lxml']
    for package in packages:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', package, '-q'])

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import markdown
    from bs4 import BeautifulSoup
except ImportError:
    print("Đang cài đặt các thư viện cần thiết...")
    install_packages()
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import markdown
    from bs4 import BeautifulSoup

import os
import re

def read_markdown_file(filepath):
    """Đọc file markdown"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def create_api_document():
    """Tạo document Word API Documentation"""
    doc = Document()
    
    # Thiết lập margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ===== TRANG BÌA =====
    # Tiêu đề chính
    title = doc.add_heading('', level=0)
    title_run = title.add_run('MAGIC ENGLISH')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('API DOCUMENTATION')
    subtitle_run.font.size = Pt(28)
    subtitle_run.font.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Thông tin version
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('RESTful API Documentation\n').font.size = Pt(14)
    info.add_run('cho Ứng Dụng Học Tiếng Anh Thông Minh\n\n').font.size = Pt(14)
    info.add_run('Version: 1.0.0\n').font.size = Pt(12)
    info.add_run('Spring Boot: 3.5.8 | Java: 21\n').font.size = Pt(12)
    info.add_run('Cập nhật: January 2025').font.size = Pt(12)
    
    doc.add_page_break()
    
    # ===== MỤC LỤC =====
    doc.add_heading('MỤC LỤC', level=1)
    toc_items = [
        ('1. Tổng Quan', 2),
        ('2. Thông Tin Kết Nối', 2),
        ('3. Authentication API', 2),
        ('4. User Management API', 2),
        ('5. Vocabulary API', 2),
        ('6. Grammar API', 2),
        ('7. IELTS Test API', 2),
        ('8. TOEIC Test API', 2),
        ('9. Achievement API', 2),
        ('10. Statistics API', 2),
        ('11. File Management API', 2),
        ('12. Pronunciation API', 2),
        ('13. Audio/TTS API', 2),
        ('14. Mã Lỗi (Error Codes)', 2),
    ]
    for item, indent in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5 * (indent - 1))
    
    doc.add_page_break()
    
    # ===== 1. TỔNG QUAN =====
    doc.add_heading('1. TỔNG QUAN', level=1)
    
    doc.add_heading('1.1 Mô tả', level=2)
    doc.add_paragraph(
        'Magic English API là một RESTful API được xây dựng trên nền tảng Spring Boot 3.5.8, '
        'cung cấp các dịch vụ backend cho ứng dụng học tiếng Anh thông minh tích hợp AI.'
    )
    
    doc.add_heading('1.2 Công nghệ sử dụng', level=2)
    tech_table = doc.add_table(rows=6, cols=3)
    tech_table.style = 'Table Grid'
    tech_headers = ['Công nghệ', 'Phiên bản', 'Mô tả']
    tech_data = [
        ['Spring Boot', '3.5.8', 'Framework chính'],
        ['Spring Security', '6.x', 'Bảo mật & JWT'],
        ['Spring AI', '1.x', 'Tích hợp AI/OpenAI'],
        ['MySQL', '8.0', 'Database'],
        ['JPA/Hibernate', '-', 'ORM'],
    ]
    for i, header in enumerate(tech_headers):
        tech_table.rows[0].cells[i].text = header
        tech_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(tech_data):
        for j, cell_data in enumerate(row_data):
            tech_table.rows[i + 1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('1.3 Định dạng Response tiêu chuẩn', level=2)
    doc.add_paragraph(
        'Tất cả API response đều được bọc tự động bởi FormatRestResponse (ResponseBodyAdvice).'
    )
    
    # Response format table
    resp_table = doc.add_table(rows=5, cols=3)
    resp_table.style = 'Table Grid'
    resp_data = [
        ['Field', 'Type', 'Mô tả'],
        ['statusCode', 'Integer', 'HTTP status code (200, 201, 400, ...)'],
        ['error', 'String', 'Thông báo lỗi (null nếu thành công)'],
        ['message', 'String/Array', 'Message từ @ApiMessage hoặc "CALL API SUCCESS"'],
        ['data', 'Object/Array', 'Dữ liệu response thực tế'],
    ]
    for i, row_data in enumerate(resp_data):
        for j, cell_data in enumerate(row_data):
            resp_table.rows[i].cells[j].text = cell_data
            if i == 0:
                resp_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Ngoại lệ không bọc response:')
    doc.add_paragraph('• Response dạng String hoặc Resource (file download)', style='List Bullet')
    doc.add_paragraph('• Swagger/OpenAPI endpoints (/v3/api-docs, /swagger-ui)', style='List Bullet')
    doc.add_paragraph('• Response có status >= 400 (errors trả về trực tiếp)', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 2. THÔNG TIN KẾT NỐI =====
    doc.add_heading('2. THÔNG TIN KẾT NỐI', level=1)
    
    conn_table = doc.add_table(rows=4, cols=2)
    conn_table.style = 'Table Grid'
    conn_data = [
        ['Thuộc tính', 'Giá trị'],
        ['Base URL', 'http://localhost:8080/api/v1'],
        ['Content-Type', 'application/json'],
        ['Authentication', 'Bearer Token (JWT)'],
    ]
    for i, row_data in enumerate(conn_data):
        for j, cell_data in enumerate(row_data):
            conn_table.rows[i].cells[j].text = cell_data
            if i == 0:
                conn_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Headers bắt buộc cho các API yêu cầu xác thực:')
    p = doc.add_paragraph()
    p.add_run('Authorization: Bearer <access_token>\n').font.name = 'Consolas'
    p.add_run('Content-Type: application/json').font.name = 'Consolas'
    
    doc.add_page_break()
    
    # ===== 3. AUTHENTICATION API =====
    doc.add_heading('3. AUTHENTICATION API', level=1)
    
    # 3.1 Login
    add_api_section(doc, 
        title='3.1 Đăng nhập (Login)',
        method='POST',
        endpoint='/api/v1/auth/login',
        auth=False,
        description='Xác thực người dùng và trả về token',
        request_body='{\n  "Email": "user@example.com",\n  "password": "password123"\n}',
        response_body='{\n  "access_token": "eyJhbGciOiJIUzI1NiIs...",\n  "refresh_token": "eyJhbGciOiJIUzI1NiIs...",\n  "user": {\n    "id": 1,\n    "email": "user@example.com",\n    "name": "Nguyen Van A",\n    "avatarUrl": "..."\n  }\n}'
    )
    
    # 3.2 Register
    add_api_section(doc,
        title='3.2 Đăng ký (Register)',
        method='POST',
        endpoint='/api/v1/auth/register',
        auth=False,
        description='Tạo tài khoản người dùng mới',
        request_body='{\n  "name": "Nguyen Van A",\n  "email": "user@example.com",\n  "password": "password123"\n}',
        response_body='{\n  "email": "user@example.com",\n  "name": "Nguyen Van A"\n}'
    )
    
    # 3.3 Get Account
    add_api_section(doc,
        title='3.3 Lấy thông tin tài khoản',
        method='GET',
        endpoint='/api/v1/auth/account',
        auth=True,
        description='Lấy thông tin tài khoản đang đăng nhập',
        response_body='{\n  "user": {\n    "id": 1,\n    "email": "user@example.com",\n    "name": "Nguyen Van A"\n  }\n}'
    )
    
    # 3.4 Refresh Token
    add_api_section(doc,
        title='3.4 Làm mới Token',
        method='GET',
        endpoint='/api/v1/auth/refresh',
        auth=True,
        description='Tạo access token mới từ refresh token (trong Cookie)'
    )
    
    # 3.5 Logout
    add_api_section(doc,
        title='3.5 Đăng xuất (Logout)',
        method='POST',
        endpoint='/api/v1/auth/logout',
        auth=True,
        description='Đăng xuất và xóa refresh token'
    )
    
    doc.add_page_break()
    
    # ===== 4. USER MANAGEMENT API =====
    doc.add_heading('4. USER MANAGEMENT API', level=1)
    
    add_api_section(doc,
        title='4.1 Tạo người dùng mới',
        method='POST',
        endpoint='/api/v1/users',
        auth=True,
        description='Tạo người dùng mới'
    )
    
    add_api_section(doc,
        title='4.2 Lấy thông tin người dùng theo ID',
        method='GET',
        endpoint='/api/v1/users/{id}',
        auth=True,
        description='Lấy chi tiết thông tin user',
        path_params=[('id', 'Long', 'ID của user')]
    )
    
    add_api_section(doc,
        title='4.3 Cập nhật người dùng',
        method='PUT',
        endpoint='/api/v1/users',
        auth=True,
        description='Cập nhật thông tin user',
        request_body='{\n  "id": 1,\n  "name": "Nguyen Van B",\n  "avatarUrl": "https://..."\n}'
    )
    
    add_api_section(doc,
        title='4.4 Xóa người dùng',
        method='DELETE',
        endpoint='/api/v1/users/{id}',
        auth=True,
        description='Xóa user theo ID'
    )
    
    doc.add_page_break()
    
    # ===== 5. VOCABULARY API =====
    doc.add_heading('5. VOCABULARY API', level=1)
    doc.add_paragraph('Quản lý từ vựng với tính năng AI enrichment.')
    
    add_api_section(doc,
        title='5.1 Thêm từ vựng mới',
        method='POST',
        endpoint='/api/v1/vocabulary',
        auth=True,
        description='Thêm từ mới, tự động làm giàu dữ liệu bằng AI',
        request_body='{\n  "word": "achievement"\n}',
        response_body='{\n  "id": 1,\n  "word": "achievement",\n  "ipa": "/əˈtʃiːvmənt/",\n  "audioUrl": "...",\n  "meaning": "Thành tựu, thành tích",\n  "wordType": "noun",\n  "examples": [...],\n  "cefrLevel": "B1"\n}'
    )
    
    add_api_section(doc,
        title='5.2 Lấy danh sách từ vựng',
        method='GET',
        endpoint='/api/v1/vocabulary',
        auth=True,
        description='Lấy danh sách từ vựng với phân trang và tìm kiếm',
        query_params=[
            ('search', 'String', 'Từ khóa tìm kiếm'),
            ('page', 'Integer', 'Số trang (mặc định: 0)'),
            ('size', 'Integer', 'Số phần tử mỗi trang (mặc định: 10)')
        ]
    )
    
    add_api_section(doc,
        title='5.3 Xem trước từ vựng (Preview)',
        method='POST',
        endpoint='/api/v1/vocabulary/preview',
        auth=True,
        description='Xem trước thông tin từ vựng mà không lưu vào database'
    )
    
    doc.add_page_break()
    
    # ===== 6. GRAMMAR API =====
    doc.add_heading('6. GRAMMAR API', level=1)
    doc.add_paragraph('Kiểm tra và sửa lỗi ngữ pháp bằng AI.')
    
    add_api_section(doc,
        title='6.1 Kiểm tra ngữ pháp',
        method='POST',
        endpoint='/api/v1/grammar/check',
        auth=True,
        description='Phân tích và sửa lỗi ngữ pháp bằng AI',
        request_body='{\n  "text": "She don\'t like to swimming in the sea."\n}',
        response_body='{\n  "id": 1,\n  "inputText": "...",\n  "correctedText": "She doesn\'t like swimming in the sea.",\n  "score": 65,\n  "errors": [...]\n}'
    )
    
    add_api_section(doc,
        title='6.2 Lấy lịch sử kiểm tra ngữ pháp',
        method='GET',
        endpoint='/api/v1/grammar',
        auth=True,
        description='Lấy danh sách các lần kiểm tra ngữ pháp'
    )
    
    add_api_section(doc,
        title='6.3 Lấy chi tiết kiểm tra ngữ pháp',
        method='GET',
        endpoint='/api/v1/grammar/{id}',
        auth=True,
        description='Lấy chi tiết một lần kiểm tra ngữ pháp'
    )
    
    add_api_section(doc,
        title='6.4 Xóa bản ghi kiểm tra ngữ pháp',
        method='DELETE',
        endpoint='/api/v1/grammar/{id}',
        auth=True,
        description='Xóa bản ghi kiểm tra ngữ pháp'
    )
    
    doc.add_page_break()
    
    # ===== 7. IELTS TEST API =====
    doc.add_heading('7. IELTS TEST API', level=1)
    doc.add_paragraph('Tạo và quản lý bài thi IELTS với AI.')
    
    add_api_section(doc,
        title='7.1 Tạo bài thi IELTS mới',
        method='POST',
        endpoint='/api/v1/ielts/generate',
        auth=True,
        description='Tạo bài thi IELTS mới bằng AI',
        request_body='{\n  "skill": "Reading",\n  "level": "Academic",\n  "difficulty": "Medium"\n}'
    )
    doc.add_paragraph('Các giá trị hợp lệ:')
    doc.add_paragraph('• skill: Reading, Writing, Listening, Speaking', style='List Bullet')
    doc.add_paragraph('• level: General, Academic', style='List Bullet')
    doc.add_paragraph('• difficulty: Easy, Medium, Hard', style='List Bullet')
    
    add_api_section(doc,
        title='7.2 Lấy bài thi theo ID',
        method='GET',
        endpoint='/api/v1/ielts/tests/{testId}',
        auth=True,
        description='Lấy chi tiết bài thi theo ID'
    )
    
    add_api_section(doc,
        title='7.3 Bắt đầu làm bài thi',
        method='POST',
        endpoint='/api/v1/ielts/start',
        auth=True,
        description='Bắt đầu một phiên làm bài thi',
        request_body='{\n  "testId": 1\n}'
    )
    
    add_api_section(doc,
        title='7.4 Nộp bài thi',
        method='POST',
        endpoint='/api/v1/ielts/submit',
        auth=True,
        description='Nộp bài thi và nhận kết quả',
        request_body='{\n  "historyId": 1,\n  "answers": [\n    {"questionId": 1, "selectedAnswerId": 2}\n  ]\n}'
    )
    
    add_api_section(doc,
        title='7.5 Lấy lịch sử làm bài',
        method='GET',
        endpoint='/api/v1/ielts/history',
        auth=True,
        description='Lấy lịch sử các bài thi đã làm'
    )
    
    doc.add_page_break()
    
    # ===== 8. TOEIC TEST API =====
    doc.add_heading('8. TOEIC TEST API', level=1)
    doc.add_paragraph('Tạo và quản lý bài thi TOEIC với AI.')
    
    add_api_section(doc,
        title='8.1 Tạo bài thi TOEIC mới',
        method='POST',
        endpoint='/api/v1/toeic/generate',
        auth=True,
        description='Tạo bài thi TOEIC mới bằng AI',
        request_body='{\n  "section": "Part 5",\n  "difficulty": "Medium"\n}'
    )
    doc.add_paragraph('Các giá trị hợp lệ:')
    doc.add_paragraph('• section: Listening, Reading, Part 1-7', style='List Bullet')
    doc.add_paragraph('• difficulty: Easy, Medium, Hard', style='List Bullet')
    
    add_api_section(doc,
        title='8.2 Lấy bài thi theo ID',
        method='GET',
        endpoint='/api/v1/toeic/tests/{testId}',
        auth=True,
        description='Lấy chi tiết bài thi theo ID'
    )
    
    add_api_section(doc,
        title='8.3 Bắt đầu làm bài thi',
        method='POST',
        endpoint='/api/v1/toeic/start',
        auth=True,
        description='Bắt đầu một phiên làm bài thi'
    )
    
    add_api_section(doc,
        title='8.4 Nộp bài thi',
        method='POST',
        endpoint='/api/v1/toeic/submit',
        auth=True,
        description='Nộp bài thi và nhận kết quả'
    )
    
    add_api_section(doc,
        title='8.5 Lấy lịch sử làm bài',
        method='GET',
        endpoint='/api/v1/toeic/history',
        auth=True,
        description='Lấy lịch sử các bài thi đã làm'
    )
    
    doc.add_page_break()
    
    # ===== 9. ACHIEVEMENT API =====
    doc.add_heading('9. ACHIEVEMENT API', level=1)
    doc.add_paragraph('Quản lý hệ thống thành tích.')
    
    add_api_section(doc,
        title='9.1 Lấy tất cả thành tích',
        method='GET',
        endpoint='/api/v1/achievements',
        auth=True,
        description='Lấy danh sách tất cả thành tích trong hệ thống'
    )
    
    add_api_section(doc,
        title='9.2 Kiểm tra và cấp thành tích',
        method='POST',
        endpoint='/api/v1/user-achievements/check',
        auth=True,
        description='Kiểm tra và cấp thành tích mới dựa trên tiến độ',
        request_body='{\n  "metricType": "VOCABULARY",\n  "currentValue": 50\n}'
    )
    doc.add_paragraph('Các metricType hợp lệ: VOCABULARY, GRAMMAR_CHECKS, IELTS_TESTS, TOEIC_TESTS, STREAK_DAYS')
    
    add_api_section(doc,
        title='9.3 Lấy thành tích của người dùng',
        method='GET',
        endpoint='/api/v1/user-achievements',
        auth=True,
        description='Lấy danh sách thành tích đã đạt được'
    )
    
    add_api_section(doc,
        title='9.4 Reset thành tích (Testing)',
        method='DELETE',
        endpoint='/api/v1/user-achievements/reset',
        auth=True,
        description='Xóa tất cả thành tích của user (dùng để test)'
    )
    
    doc.add_page_break()
    
    # ===== 10. STATISTICS API =====
    doc.add_heading('10. STATISTICS API', level=1)
    doc.add_paragraph('Thống kê dữ liệu học tập.')
    
    add_api_section(doc,
        title='10.1 Thống kê từ vựng theo loại từ',
        method='GET',
        endpoint='/api/v1/vocabulary/breakdown',
        auth=True,
        description='Phân loại từ vựng theo loại từ (noun, verb, adjective, ...)'
    )
    
    add_api_section(doc,
        title='10.2 Phân bố CEFR',
        method='GET',
        endpoint='/api/v1/vocabulary/cefr-distribution',
        auth=True,
        description='Phân bố từ vựng theo cấp độ CEFR (A1-C2)'
    )
    
    add_api_section(doc,
        title='10.3 Tổng số từ vựng',
        method='GET',
        endpoint='/api/v1/vocabulary/count',
        auth=True,
        description='Lấy tổng số từ vựng đã học'
    )
    
    add_api_section(doc,
        title='10.4 Thống kê trang chủ',
        method='GET',
        endpoint='/api/v1/vocabulary/home-stats',
        auth=True,
        description='Thống kê tổng hợp cho dashboard (wordsToday, totalWords, streakDays, grammarChecks, avgGrammarScore)'
    )
    
    add_api_section(doc,
        title='10.5 Thống kê từ vựng theo ngày',
        method='GET',
        endpoint='/api/v1/stats/daily-vocabulary',
        auth=True,
        description='Số từ học theo từng ngày',
        query_params=[('days', 'Integer', 'Số ngày (mặc định: 7)')]
    )
    
    add_api_section(doc,
        title='10.6 Thống kê kiểm tra ngữ pháp theo ngày',
        method='GET',
        endpoint='/api/v1/stats/daily-grammar-checks',
        auth=True,
        description='Số lần kiểm tra ngữ pháp theo ngày'
    )
    
    add_api_section(doc,
        title='10.7 Điểm ngữ pháp theo ngày',
        method='GET',
        endpoint='/api/v1/stats/daily-grammar-scores',
        auth=True,
        description='Điểm trung bình ngữ pháp theo ngày'
    )
    
    add_api_section(doc,
        title='10.8 Hoạt động theo ngày',
        method='GET',
        endpoint='/api/v1/stats/daily-activity',
        auth=True,
        description='Kiểm tra hoạt động học tập theo ngày (streak)'
    )
    
    doc.add_page_break()
    
    # ===== 11. FILE MANAGEMENT API =====
    doc.add_heading('11. FILE MANAGEMENT API', level=1)
    doc.add_paragraph('Upload và download files.')
    
    add_api_section(doc,
        title='11.1 Upload file',
        method='POST',
        endpoint='/api/v1/files',
        auth=True,
        description='Upload file (ảnh, document). Content-Type: multipart/form-data'
    )
    doc.add_paragraph('Form Data: file (File), folder (String)')
    doc.add_paragraph('Allowed Extensions: pdf, jpg, jpeg, png, doc, docx')
    
    add_api_section(doc,
        title='11.2 Download file',
        method='GET',
        endpoint='/api/v1/files',
        auth=True,
        description='Tải file xuống',
        query_params=[
            ('fileName', 'String', 'Tên file'),
            ('folder', 'String', 'Thư mục chứa file')
        ]
    )
    
    doc.add_page_break()
    
    # ===== 12. PRONUNCIATION API =====
    doc.add_heading('12. PRONUNCIATION API', level=1)
    doc.add_paragraph('Phân tích và đánh giá phát âm.')
    
    add_api_section(doc,
        title='12.1 Phân tích phát âm',
        method='POST',
        endpoint='/api/v1/pronunciation/analyze',
        auth=True,
        description='Đánh giá phát âm dựa trên speech-to-text',
        request_body='{\n  "expectedWord": "achievement",\n  "transcribedText": "achievment",\n  "ipa": "/əˈtʃiːvmənt/"\n}',
        response_body='{\n  "score": 75,\n  "isCorrect": false,\n  "expectedWord": "achievement",\n  "transcribedWord": "achievment",\n  "feedback": "...",\n  "suggestions": [...]\n}'
    )
    
    doc.add_page_break()
    
    # ===== 13. AUDIO/TTS API =====
    doc.add_heading('13. AUDIO/TTS API', level=1)
    doc.add_paragraph('Proxy audio và Text-to-Speech.')
    
    add_api_section(doc,
        title='13.1 Proxy Audio',
        method='GET',
        endpoint='/api/audio/proxy',
        auth=False,
        description='Proxy audio từ external URLs (giải quyết CORS)',
        query_params=[('url', 'String', 'URL audio cần proxy')]
    )
    
    add_api_section(doc,
        title='13.2 Text-to-Speech',
        method='GET',
        endpoint='/api/audio/tts',
        auth=False,
        description='Tạo audio từ text bằng VoiceRSS API',
        query_params=[('text', 'String', 'Văn bản cần đọc')]
    )
    doc.add_paragraph('Multi-voice: Sử dụng "Man:" hoặc "Woman:" để đổi giọng')
    
    add_api_section(doc,
        title='13.3 Test Audio',
        method='GET',
        endpoint='/api/audio/test',
        auth=False,
        description='Kiểm tra audio proxy hoạt động'
    )
    
    doc.add_page_break()
    
    # ===== 14. MÃ LỖI =====
    doc.add_heading('14. MÃ LỖI (ERROR CODES)', level=1)
    
    doc.add_heading('14.1 HTTP Status Codes', level=2)
    error_table = doc.add_table(rows=9, cols=3)
    error_table.style = 'Table Grid'
    error_data = [
        ['Code', 'Status', 'Mô tả'],
        ['200', 'OK', 'Thành công'],
        ['201', 'Created', 'Tạo thành công'],
        ['204', 'No Content', 'Xóa thành công'],
        ['400', 'Bad Request', 'Request không hợp lệ'],
        ['401', 'Unauthorized', 'Chưa xác thực'],
        ['403', 'Forbidden', 'Không có quyền'],
        ['404', 'Not Found', 'Không tìm thấy'],
        ['500', 'Internal Server Error', 'Lỗi server'],
    ]
    for i, row_data in enumerate(error_data):
        for j, cell_data in enumerate(row_data):
            error_table.rows[i].cells[j].text = cell_data
            if i == 0:
                error_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_heading('14.2 Common Errors', level=2)
    doc.add_paragraph('• IdInvalidException: ID không hợp lệ hoặc không tồn tại', style='List Bullet')
    doc.add_paragraph('• StorageException: Lỗi upload file (extension, size)', style='List Bullet')
    doc.add_paragraph('• Token expired: Token hết hạn - sử dụng refresh token', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('─' * 50 + '\n')
    footer.add_run('Magic English Team\n').font.bold = True
    footer.add_run('Version 1.0.0 | January 2025\n')
    footer.add_run('Made with ❤️ by Magic English Team')
    
    return doc


def add_api_section(doc, title, method, endpoint, auth, description, 
                    request_body=None, response_body=None, 
                    path_params=None, query_params=None):
    """Thêm một section API vào document"""
    doc.add_heading(title, level=2)
    
    # Info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_table.rows[0].cells[0].text = 'Endpoint'
    info_table.rows[0].cells[1].text = f'{method} {endpoint}'
    info_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    
    info_table.rows[1].cells[0].text = 'Method'
    info_table.rows[1].cells[1].text = method
    info_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    
    info_table.rows[2].cells[0].text = 'Auth'
    info_table.rows[2].cells[1].text = '✅ Yêu cầu Bearer Token' if auth else '❌ Không yêu cầu'
    info_table.rows[2].cells[0].paragraphs[0].runs[0].bold = True
    
    info_table.rows[3].cells[0].text = 'Description'
    info_table.rows[3].cells[1].text = description
    info_table.rows[3].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Path Parameters
    if path_params:
        doc.add_paragraph('Path Parameters:', style='Intense Emphasis')
        param_table = doc.add_table(rows=len(path_params) + 1, cols=3)
        param_table.style = 'Table Grid'
        param_table.rows[0].cells[0].text = 'Parameter'
        param_table.rows[0].cells[1].text = 'Type'
        param_table.rows[0].cells[2].text = 'Description'
        for i, (name, ptype, desc) in enumerate(path_params):
            param_table.rows[i + 1].cells[0].text = name
            param_table.rows[i + 1].cells[1].text = ptype
            param_table.rows[i + 1].cells[2].text = desc
        doc.add_paragraph()
    
    # Query Parameters
    if query_params:
        doc.add_paragraph('Query Parameters:', style='Intense Emphasis')
        param_table = doc.add_table(rows=len(query_params) + 1, cols=3)
        param_table.style = 'Table Grid'
        param_table.rows[0].cells[0].text = 'Parameter'
        param_table.rows[0].cells[1].text = 'Type'
        param_table.rows[0].cells[2].text = 'Description'
        for i, (name, ptype, desc) in enumerate(query_params):
            param_table.rows[i + 1].cells[0].text = name
            param_table.rows[i + 1].cells[1].text = ptype
            param_table.rows[i + 1].cells[2].text = desc
        doc.add_paragraph()
    
    # Request Body
    if request_body:
        doc.add_paragraph('Request Body:', style='Intense Emphasis')
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(request_body)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9)
        doc.add_paragraph()
    
    # Response Body
    if response_body:
        doc.add_paragraph('Response:', style='Intense Emphasis')
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(response_body)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9)
        doc.add_paragraph()


if __name__ == '__main__':
    print("=" * 50)
    print("MAGIC ENGLISH - API Documentation Export")
    print("=" * 50)
    
    # Tạo document
    print("\n[1/2] Đang tạo document Word...")
    doc = create_api_document()
    
    # Lưu file
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'API_Documentation.docx')
    print(f"[2/2] Đang lưu file: {output_path}")
    doc.save(output_path)
    
    print("\n" + "=" * 50)
    print("✅ HOÀN TẤT!")
    print(f"📄 File đã được tạo: API_Documentation.docx")
    print("=" * 50)
